#!/usr/bin/env python3
"""Generate development cost estimate (смета) Word document for Star Carpet."""

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor

OUT = Path(__file__).resolve().parents[1] / "Смета-STAR-CARPET-разработка.docx"

# Курс ЦБ РФ на 20.07.2026
USD_RATE = 78.3987

BURGUNDY = RGBColor(0x7A, 0x24, 0x30)
INK = RGBColor(0x1A, 0x17, 0x14)
MUTED = RGBColor(0x5C, 0x56, 0x4E)
GREEN = RGBColor(0x1F, 0x6B, 0x3A)


def rub_usd(rub: int) -> str:
    usd = rub / USD_RATE
    return f"{rub:,} ₽ / ≈ ${usd:,.0f}".replace(",", " ")


def set_run(run, *, bold=False, size=11, color=INK, font="Calibri"):
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.name = font
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), font)
    rFonts.set(qn("w:hAnsi"), font)
    rFonts.set(qn("w:eastAsia"), font)
    rFonts.set(qn("w:cs"), font)
    rPr.insert(0, rFonts)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    run = p.add_run(text)
    if level == 0:
        set_run(run, bold=True, size=22, color=BURGUNDY)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif level == 1:
        set_run(run, bold=True, size=16, color=BURGUNDY)
        p.space_before = Pt(18)
        p.space_after = Pt(8)
    else:
        set_run(run, bold=True, size=13, color=INK)
        p.space_before = Pt(14)
        p.space_after = Pt(6)
    return p


def add_p(doc, text, *, bold=False, size=11, color=INK, center=False, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run(run, bold=bold, size=size, color=color)
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.15
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        set_run(run, size=11)
        p.paragraph_format.space_after = Pt(3)


def shade_cell(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def set_cell_text(cell, text, *, bold=False, size=10, color=INK):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    set_run(run, bold=bold, size=size, color=color)


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_text(cell, h, bold=True, size=10, color=RGBColor(0xFF, 0xFF, 0xFF))
        shade_cell(cell, "7A2430")
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            set_cell_text(cell, val, size=10)
            if r_idx % 2 == 1:
                shade_cell(cell, "F6F3EE")
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return table


def main():
    site = 130_000
    crm = 170_000
    total = site + crm  # 300_000
    friend_price = 210_000
    discount = total - friend_price  # 90_000
    prepay_full = round(friend_price * 0.93)  # 195_300
    half = friend_price // 2  # 105_000

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    add_heading(doc, "СМЕТА НА РАЗРАБОТКУ", 0)
    add_heading(doc, "Сайт + личная CRM", 0)
    add_p(
        doc,
        "Star Carpet — комплексная разработка под ключ",
        center=True,
        color=MUTED,
        size=11,
    )

    add_table(
        doc,
        ["Параметр", "Значение"],
        [
            ["Заказчик", "Star Carpet"],
            ["Продукт", "Сайт + личная CRM-система"],
            ["Срок работ", "1 месяц"],
            ["Курс USD (ЦБ РФ на 20.07.2026)", f"{USD_RATE:.4f} ₽"],
            ["Дата сметы", "20 июля 2026"],
        ],
        col_widths=[6, 11],
    )

    # 1. Состав работ и цены
    add_heading(doc, "1. Состав работ и стоимость", 1)
    add_table(
        doc,
        ["№", "Позиция", "Стоимость, ₽", "≈ USD"],
        [
            ["1", "Разработка сайта", f"{site:,}".replace(",", " "), f"≈ ${site / USD_RATE:,.0f}".replace(",", " ")],
            [
                "2",
                "Личная CRM-система под сайт",
                f"{crm:,}".replace(",", " "),
                f"≈ ${crm / USD_RATE:,.0f}".replace(",", " "),
            ],
            [
                "",
                "Итого по прайсу",
                f"{total:,}".replace(",", " "),
                f"≈ ${total / USD_RATE:,.0f}".replace(",", " "),
            ],
        ],
        col_widths=[1.5, 8, 4, 3.5],
    )

    add_p(doc, "Что входит в позиции:", bold=True, space_after=2)
    add_bullets(
        doc,
        [
            "Сайт — дизайн, вёрстка, каталог, карточки товаров, адаптив, базовый SEO, запуск.",
            "Личная CRM — система под процессы магазина: заявки, клиенты, статусы, связь с сайтом.",
        ],
    )

    # 2. Скидка
    add_heading(doc, "2. Скидка для знакомых", 1)
    add_p(
        doc,
        f"Прайс-стоимость пакета: {rub_usd(total)}",
    )
    add_p(
        doc,
        f"Скидка для знакомых: −{discount:,} ₽ (≈ −${discount / USD_RATE:,.0f})".replace(",", " "),
        bold=True,
    )
    add_p(
        doc,
        f"Итоговая стоимость: {rub_usd(friend_price)}",
        bold=True,
        size=14,
        color=BURGUNDY,
    )

    add_table(
        doc,
        ["Показатель", "Сумма"],
        [
            ["Прайс (сайт + CRM)", rub_usd(total)],
            ["Скидка для знакомых", f"−{discount:,} ₽ / ≈ −${discount / USD_RATE:,.0f}".replace(",", " ")],
            ["К оплате", rub_usd(friend_price)],
        ],
        col_widths=[8, 9],
    )

    # 3. Оплата
    add_heading(doc, "3. Варианты оплаты", 1)

    add_heading(doc, "Вариант A — стандарт (рекомендуемый)", 2)
    add_p(doc, "50% до начала работы + 50% во время работы (в течение 1–2 недель).", bold=True)
    add_table(
        doc,
        ["Этап", "Срок", "Сумма"],
        [
            [
                "Предоплата 50%",
                "До старта работ",
                rub_usd(half),
            ],
            [
                "Второй платёж 50%",
                "В течение 1–2 недель работы",
                rub_usd(half),
            ],
            ["Итого", "", rub_usd(friend_price)],
        ],
        col_widths=[5.5, 5.5, 6],
    )

    add_heading(doc, "Вариант B — 100% предоплата", 2)
    add_p(
        doc,
        "При полной предоплате — дополнительная скидка 7%.",
        bold=True,
    )
    add_table(
        doc,
        ["Параметр", "Значение"],
        [
            ["Базовая цена со скидкой для знакомых", rub_usd(friend_price)],
            ["Скидка за 100% предоплату (−7%)", f"−{friend_price - prepay_full:,} ₽ / ≈ −${(friend_price - prepay_full) / USD_RATE:,.0f}".replace(",", " ")],
            ["Итого при 100% предоплате", rub_usd(prepay_full)],
            ["Оплата", "100% до начала работ"],
        ],
        col_widths=[9, 8],
    )
    add_p(
        doc,
        f"Выгода варианта B: экономия {friend_price - prepay_full:,} ₽ "
        f"(≈ ${(friend_price - prepay_full) / USD_RATE:,.0f}) относительно варианта A.".replace(",", " "),
        color=GREEN,
        bold=True,
    )

    # 4. Сроки
    add_heading(doc, "4. Сроки работ", 1)
    add_p(doc, "Срок выполнения: 1 месяц.", bold=True, size=13)
    add_p(
        doc,
        "Срок считается от даты предоплаты и предоставления доступов/материалов. "
        "Оперативные согласования ускоряют сдачу внутри заявленного месяца.",
        color=MUTED,
    )

    # 5. Поддержка
    add_heading(doc, "5. Поддержка на год", 1)
    add_p(
        doc,
        "Техническое и дизайн-обслуживание на 12 месяцев предоставляется бесплатно "
        "при выполнении обоих условий:",
        bold=True,
    )
    add_bullets(
        doc,
        [
            "оплата проекта предоплатой (вариант B — 100% предоплата);",
            "гарантия рекомендации / рекламы нас другим людям (устные или письменные рекомендации клиентам и партнёрам).",
        ],
    )
    add_p(doc, "Что входит в годовую поддержку:", bold=True, space_after=2)
    add_bullets(
        doc,
        [
            "техническое обслуживание сайта и CRM (исправление ошибок, мелкие правки, мониторинг работоспособности);",
            "дизайн-обслуживание (небольшие правки макетов, баннеров, блоков в рамках текущего стиля);",
            "консультации по использованию системы.",
        ],
    )
    add_p(
        doc,
        "Крупные новые функции и переработки сверх текущего объёма — отдельно по согласованию.",
        color=MUTED,
        size=10,
    )

    # 6. Итог
    add_heading(doc, "6. Итоговая сводка", 1)
    add_table(
        doc,
        ["Параметр", "Вариант A (50/50)", "Вариант B (100% предоплата)"],
        [
            ["Стоимость", rub_usd(friend_price), rub_usd(prepay_full)],
            [
                "Оплата",
                f"{rub_usd(half)} старт + {rub_usd(half)} за 1–2 нед.",
                "100% до старта",
            ],
            ["Срок", "1 месяц", "1 месяц"],
            [
                "Поддержка 12 мес. (тех + дизайн)",
                "По согласованию",
                "Бесплатно* при рекомендации",
            ],
        ],
        col_widths=[5, 6, 6],
    )
    add_p(
        doc,
        "*Бесплатная поддержка на год — при 100% предоплате и гарантии рекламы/рекомендаций другим людям.",
        color=MUTED,
        size=10,
    )

    add_heading(doc, "ИТОГО К ОПЛАТЕ", 1)
    add_p(doc, f"Прайс: {rub_usd(total)}", center=True, size=12)
    add_p(
        doc,
        f"Со скидкой для знакомых: {rub_usd(friend_price)}",
        center=True,
        bold=True,
        size=16,
        color=BURGUNDY,
    )
    add_p(
        doc,
        f"При 100% предоплате (−7%): {rub_usd(prepay_full)}",
        center=True,
        bold=True,
        size=14,
        color=GREEN,
    )
    add_p(doc, "Срок: 1 месяц", center=True, bold=True, size=12)
    add_p(
        doc,
        f"Курс для пересчёта в USD: {USD_RATE:.4f} ₽ (ЦБ РФ, 20.07.2026). "
        "Суммы в долларах — ориентировочные, округлены.",
        center=True,
        color=MUTED,
        size=9,
    )

    doc.save(OUT)
    print(f"Saved: {OUT}")


if __name__ == "__main__":
    main()
